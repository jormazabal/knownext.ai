from __future__ import annotations

import html
import re
import tempfile
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Iterable

from fastapi import HTTPException
from markdown_it import MarkdownIt
from mdit_py_plugins.tasklists import tasklists_plugin

from app.schemas.config import ExportTemplateConfig
from app.schemas.document import DocumentExportFormat, ExportDocumentResponse
from app.services.export_template_service import export_template_service
from app.services.filesystem_service import decode_document_id
from app.services.project_service import project_service


@dataclass
class InlineRun:
    text: str = ""
    bold: bool = False
    italic: bool = False
    strike: bool = False
    underline: bool = False
    code: bool = False
    link: str | None = None
    image: str | None = None
    alt: str = ""


@dataclass
class ExportBlock:
    kind: str
    inlines: list[InlineRun] | None = None
    level: int = 0
    rows: list[list[list[InlineRun]]] | None = None
    ordered: bool = False
    prefix: str = ""
    quote: bool = False
    text: str = ""


_IMAGE_TARGET_PATTERN = re.compile(r"^(https?:|data:|mailto:|#)", re.IGNORECASE)
_UNDERLINE_INLINE_PATTERN = re.compile(r"^<u>(.*?)</u>$", re.IGNORECASE | re.DOTALL)


def _now_iso() -> str:
    return datetime.now(timezone.utc).isoformat()


def _mm_to_points(value: float) -> float:
    return float(value) * 72 / 25.4


def _hex_to_rgb_tuple(color: str) -> tuple[int, int, int]:
    value = color.lstrip("#")
    return int(value[0:2], 16), int(value[2:4], 16), int(value[4:6], 16)


def _docx_color(color: str):
    from docx.shared import RGBColor

    return RGBColor(*_hex_to_rgb_tuple(color))


def _reportlab_color(color: str):
    from reportlab.lib.colors import HexColor

    return HexColor(color)


def _pdf_font_name(font_family: str, *, bold: bool = False, italic: bool = False) -> str:
    normalized = font_family.lower()
    if "courier" in normalized or "consolas" in normalized or "mono" in normalized:
        if bold and italic:
            return "Courier-BoldOblique"
        if bold:
            return "Courier-Bold"
        if italic:
            return "Courier-Oblique"
        return "Courier"
    if "times" in normalized or "georgia" in normalized:
        if bold and italic:
            return "Times-BoldItalic"
        if bold:
            return "Times-Bold"
        if italic:
            return "Times-Italic"
        return "Times-Roman"
    if bold and italic:
        return "Helvetica-BoldOblique"
    if bold:
        return "Helvetica-Bold"
    if italic:
        return "Helvetica-Oblique"
    return "Helvetica"


def _read_style(template: ExportTemplateConfig, block: ExportBlock) -> dict:
    if block.kind == "heading":
        style = template.headings.get(f"h{block.level}") or template.normal
        return style.model_dump() if hasattr(style, "model_dump") else dict(style)
    if block.kind == "code":
        return template.code.model_dump()
    return template.normal.model_dump()


def _style_is_bold(style: dict) -> bool:
    return style.get("textFormat") in {"bold", "bold_underline"}


def _style_is_underlined(style: dict) -> bool:
    return style.get("textFormat") in {"underline", "bold_underline"}


def _sanitize_export_path(output_path: str, export_format: DocumentExportFormat) -> Path:
    raw_path = Path(output_path).expanduser()
    if not str(raw_path).strip():
        raise HTTPException(status_code=400, detail="Output path is required")
    suffix = f".{export_format}"
    if raw_path.suffix.lower() != suffix:
        raw_path = raw_path.with_suffix(suffix)
    if raw_path.parent and not raw_path.parent.exists():
        raise HTTPException(status_code=400, detail="Output folder does not exist")
    if raw_path.exists() and raw_path.is_dir():
        raise HTTPException(status_code=400, detail="Output path points to a folder")
    return raw_path


def _document_context(document_id: str) -> tuple[Path, str]:
    if not document_id.startswith("fs_"):
        raise HTTPException(status_code=404, detail="Document not found")
    project_id, relative_path = decode_document_id(document_id)
    return project_service._get_project_root(project_id), relative_path


def _resolve_image_path(project_root: Path, document_path: str, target: str) -> Path | None:
    if not target or _IMAGE_TARGET_PATTERN.match(target):
        return None
    clean_target = target.split("#", 1)[0].split("?", 1)[0]
    try:
        clean_target = clean_target.replace("%20", " ")
    except Exception:
        pass
    candidate = Path(clean_target)
    if not candidate.is_absolute():
        candidate = project_root / Path(document_path).parent / clean_target
    try:
        resolved = candidate.resolve()
        resolved.relative_to(project_root.resolve())
    except ValueError:
        return None
    return resolved if resolved.is_file() else None


def _markdown_parser() -> MarkdownIt:
    return MarkdownIt("gfm-like", {"html": True, "linkify": False}).use(tasklists_plugin, enabled=True)


def _inline_runs(children: Iterable, state: dict | None = None) -> list[InlineRun]:
    runs: list[InlineRun] = []
    active = dict(state or {})
    children = list(children or [])
    index = 0
    while index < len(children):
        token = children[index]
        token_type = token.type
        if token_type == "text":
            if token.content:
                runs.append(InlineRun(token.content, **active))
        elif token_type in {"softbreak", "hardbreak"}:
            runs.append(InlineRun("\n", **active))
        elif token_type == "code_inline":
            runs.append(InlineRun(token.content, **{**active, "code": True}))
        elif token_type == "image":
            runs.append(InlineRun(image=token.attrGet("src"), alt=token.content or token.attrGet("alt") or "", **active))
        elif token_type == "html_inline":
            content = token.content.strip()
            match = _UNDERLINE_INLINE_PATTERN.match(content)
            if match:
                runs.append(InlineRun(match.group(1), **{**active, "underline": True}))
            elif "task-list-item-checkbox" in content:
                runs.append(InlineRun("[x] " if "checked" in content else "[ ] ", **active))
            elif content.lower() == "<u>":
                active["underline"] = True
            elif content.lower() == "</u>":
                active["underline"] = False
            elif content:
                runs.append(InlineRun(content, **active))
        elif token_type.endswith("_open"):
            mark = token_type.removesuffix("_open")
            next_state = dict(active)
            if mark == "strong":
                next_state["bold"] = True
            elif mark == "em":
                next_state["italic"] = True
            elif mark == "s":
                next_state["strike"] = True
            elif mark == "link":
                next_state["link"] = token.attrGet("href")
            depth = 1
            nested = []
            close_type = f"{mark}_close"
            index += 1
            while index < len(children):
                current = children[index]
                if current.type == token_type:
                    depth += 1
                if current.type == close_type:
                    depth -= 1
                    if depth == 0:
                        break
                nested.append(current)
                index += 1
            runs.extend(_inline_runs(nested, next_state))
        index += 1
    return runs


def _parse_table(tokens: list, start: int) -> tuple[ExportBlock, int]:
    rows: list[list[list[InlineRun]]] = []
    current_row: list[list[InlineRun]] | None = None
    index = start + 1
    while index < len(tokens):
        token = tokens[index]
        if token.type == "table_close":
            return ExportBlock(kind="table", rows=rows), index
        if token.type == "tr_open":
            current_row = []
        elif token.type == "tr_close" and current_row is not None:
            rows.append(current_row)
            current_row = None
        elif token.type in {"th_open", "td_open"} and current_row is not None:
            inline = tokens[index + 1] if index + 1 < len(tokens) and tokens[index + 1].type == "inline" else None
            current_row.append(_inline_runs(inline.children if inline else []))
        index += 1
    return ExportBlock(kind="table", rows=rows), index


def _parse_blocks(markdown: str) -> list[ExportBlock]:
    tokens = _markdown_parser().parse(markdown)
    blocks: list[ExportBlock] = []
    list_stack: list[dict] = []
    list_item_stack: list[dict] = []
    quote_depth = 0
    last_block_end_line: int | None = None
    index = 0

    def append_source_spacing(token) -> None:
        nonlocal last_block_end_line
        if not token.map:
            return
        start_line, end_line = token.map
        if last_block_end_line is not None:
            for _ in range(max(0, start_line - last_block_end_line - 1)):
                blocks.append(ExportBlock(kind="blank_line"))
        last_block_end_line = end_line

    while index < len(tokens):
        token = tokens[index]
        token_type = token.type
        if token_type == "heading_open":
            append_source_spacing(token)
            inline = tokens[index + 1] if index + 1 < len(tokens) and tokens[index + 1].type == "inline" else None
            blocks.append(ExportBlock(kind="heading", level=int(token.tag[1]), inlines=_inline_runs(inline.children if inline else []), quote=quote_depth > 0))
            index += 2
        elif token_type == "paragraph_open":
            append_source_spacing(token)
            inline = tokens[index + 1] if index + 1 < len(tokens) and tokens[index + 1].type == "inline" else None
            inlines = _inline_runs(inline.children if inline else [])
            if list_item_stack:
                current_item = list_item_stack[-1]
                prefix = current_item["prefix"] if not current_item["emitted"] else ""
                current_item["emitted"] = True
                blocks.append(ExportBlock(kind="list_item", inlines=inlines, prefix=prefix, level=len(list_stack), quote=quote_depth > 0))
            else:
                blocks.append(ExportBlock(kind="paragraph", inlines=inlines, quote=quote_depth > 0))
            index += 2
        elif token_type in {"fence", "code_block"}:
            append_source_spacing(token)
            blocks.append(ExportBlock(kind="code", text=token.content))
        elif token_type == "hr":
            append_source_spacing(token)
            blocks.append(ExportBlock(kind="hr"))
        elif token_type == "blockquote_open":
            quote_depth += 1
        elif token_type == "blockquote_close":
            quote_depth = max(0, quote_depth - 1)
        elif token_type == "bullet_list_open":
            list_stack.append({"type": "bullet", "counter": 0})
        elif token_type == "ordered_list_open":
            start = int(token.attrGet("start") or 1)
            list_stack.append({"type": "ordered", "counter": start - 1})
        elif token_type in {"bullet_list_close", "ordered_list_close"}:
            if list_stack:
                list_stack.pop()
        elif token_type == "list_item_open":
            current_list = list_stack[-1] if list_stack else {"type": "bullet", "counter": 0}
            current_list["counter"] += 1
            prefix = f"{current_list['counter']}. " if current_list["type"] == "ordered" else "- "
            list_item_stack.append({"prefix": prefix, "emitted": False})
        elif token_type == "list_item_close":
            if list_item_stack:
                list_item_stack.pop()
        elif token_type == "table_open":
            append_source_spacing(token)
            block, index = _parse_table(tokens, index)
            blocks.append(block)
        index += 1
    return blocks


def _plain_text(inlines: list[InlineRun] | None) -> str:
    parts = []
    for run in inlines or []:
        if run.image:
            parts.append(run.alt or run.image)
        else:
            parts.append(run.text)
    return "".join(parts)


def _default_export_filename(document_name: str, export_format: DocumentExportFormat) -> str:
    base_name = Path(document_name or "documento.md").name
    without_markdown_extension = re.sub(r"\.[Mm][Dd]$", "", base_name)
    suffix = f".{export_format}"
    if without_markdown_extension.lower().endswith(suffix):
        return without_markdown_extension
    return f"{without_markdown_extension}{suffix}"


def _blank_line_height(template: ExportTemplateConfig) -> float:
    return template.normal.fontSizePt * template.paragraph.lineSpacing


class ExportService:
    def export_document_bytes(
        self,
        document_id: str,
        export_format: DocumentExportFormat,
        markdown: str,
        document_name: str,
    ) -> bytes:
        with tempfile.TemporaryDirectory(prefix="knownext-export-") as temp_dir:
            target = Path(temp_dir) / _default_export_filename(document_name, export_format)
            self.export_document(document_id, export_format, str(target), markdown, document_name)
            return target.read_bytes()

    def export_document(
        self,
        document_id: str,
        export_format: DocumentExportFormat,
        output_path: str,
        markdown: str,
        document_name: str,
    ) -> ExportDocumentResponse:
        target = _sanitize_export_path(output_path, export_format)
        project_root, relative_path = _document_context(document_id)
        template = export_template_service.get_template()
        if export_format == "md":
            target.write_text(markdown, encoding="utf-8")
        elif export_format == "docx":
            self._write_docx(target, markdown, document_name, project_root, relative_path, template)
        elif export_format == "pdf":
            self._write_pdf(target, markdown, document_name, project_root, relative_path, template)
        else:
            raise HTTPException(status_code=400, detail="Unsupported export format")
        return ExportDocumentResponse(
            documentId=document_id,
            format=export_format,
            outputPath=str(target),
            exportedAt=_now_iso(),
        )

    def _write_docx(
        self,
        target: Path,
        markdown: str,
        document_name: str,
        project_root: Path,
        relative_path: str,
        template: ExportTemplateConfig,
    ) -> None:
        from docx import Document as DocxDocument
        from docx.enum.text import WD_LINE_SPACING
        from docx.shared import Cm, Pt

        document = DocxDocument()
        section = document.sections[0]
        section.page_width = Cm(21.59 if template.page.size == "Letter" else 21)
        section.page_height = Cm(27.94 if template.page.size == "Letter" else 29.7)
        section.top_margin = Cm(template.page.margins.topMm / 10)
        section.right_margin = Cm(template.page.margins.rightMm / 10)
        section.bottom_margin = Cm(template.page.margins.bottomMm / 10)
        section.left_margin = Cm(template.page.margins.leftMm / 10)

        for block in _parse_blocks(markdown):
            if block.kind == "heading":
                paragraph = document.add_paragraph()
                paragraph.style = document.styles[f"Heading {min(max(block.level, 1), 6)}"]
                self._format_docx_paragraph(paragraph, template, block)
                self._append_docx_runs(paragraph, block.inlines or [], template, block, project_root, relative_path)
            elif block.kind in {"paragraph", "list_item"}:
                paragraph = document.add_paragraph()
                self._format_docx_paragraph(paragraph, template, block)
                if block.kind == "list_item":
                    paragraph.paragraph_format.left_indent = Pt(16 * max(block.level, 1))
                    paragraph.add_run(block.prefix)
                self._append_docx_runs(paragraph, block.inlines or [], template, block, project_root, relative_path)
            elif block.kind == "code":
                paragraph = document.add_paragraph()
                self._format_docx_paragraph(paragraph, template, block)
                run = paragraph.add_run(block.text.rstrip("\n"))
                run.font.name = template.code.fontFamily
                run.font.size = Pt(template.code.fontSizePt)
                run.font.color.rgb = _docx_color(template.code.color)
            elif block.kind == "hr":
                paragraph = document.add_paragraph()
                paragraph.add_run("―" * 36)
            elif block.kind == "table":
                self._append_docx_table(document, block, template, project_root, relative_path)
            elif block.kind == "blank_line":
                self._append_docx_blank_line(document, template)

        document.save(target)

    def _format_docx_paragraph(self, paragraph, template: ExportTemplateConfig, block: ExportBlock) -> None:
        from docx.shared import Pt

        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(template.paragraph.spaceAfterPt)
        paragraph.paragraph_format.line_spacing = template.paragraph.lineSpacing
        if block.quote:
            paragraph.paragraph_format.left_indent = Pt(18)

    def _append_docx_blank_line(self, document, template: ExportTemplateConfig) -> None:
        from docx.shared import Pt

        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1
        run = paragraph.add_run(" ")
        run.font.name = template.normal.fontFamily
        run.font.size = Pt(template.normal.fontSizePt)

    def _append_docx_runs(self, paragraph, runs: list[InlineRun], template: ExportTemplateConfig, block: ExportBlock, project_root: Path, relative_path: str) -> None:
        from docx.shared import Inches, Pt

        style = _read_style(template, block)
        for inline in runs:
            if inline.image:
                image_path = _resolve_image_path(project_root, relative_path, inline.image)
                if image_path is not None:
                    try:
                        paragraph.add_run().add_picture(str(image_path), width=Inches(5.8))
                        continue
                    except Exception:
                        pass
                if inline.alt or inline.image:
                    paragraph.add_run(inline.alt or inline.image)
                continue
            parts = inline.text.split("\n")
            for part_index, part in enumerate(parts):
                run = paragraph.add_run(part)
                run.bold = inline.bold or (block.kind == "heading" and _style_is_bold(style))
                run.italic = inline.italic
                run.underline = inline.underline or inline.link is not None or (block.kind == "heading" and _style_is_underlined(style))
                run.font.strike = inline.strike
                run.font.name = template.code.fontFamily if inline.code else style["fontFamily"]
                run.font.size = Pt(template.code.fontSizePt if inline.code else style["fontSizePt"])
                run.font.color.rgb = _docx_color(template.document.linkColor if inline.link else template.code.color if inline.code else style["color"])
                if part_index < len(parts) - 1:
                    paragraph.add_run().add_break()

    def _append_docx_table(self, document, block: ExportBlock, template: ExportTemplateConfig, project_root: Path, relative_path: str) -> None:
        rows = block.rows or []
        if not rows:
            return
        table = document.add_table(rows=len(rows), cols=max(len(row) for row in rows))
        table.style = "Table Grid"
        for row_index, row in enumerate(rows):
            for column_index, cell_runs in enumerate(row):
                paragraph = table.rows[row_index].cells[column_index].paragraphs[0]
                cell_block = ExportBlock(kind="heading" if row_index == 0 else "paragraph", level=6)
                self._append_docx_runs(paragraph, cell_runs, template, cell_block, project_root, relative_path)

    def _write_pdf(
        self,
        target: Path,
        markdown: str,
        document_name: str,
        project_root: Path,
        relative_path: str,
        template: ExportTemplateConfig,
    ) -> None:
        from reportlab.lib.pagesizes import A4, LETTER
        from reportlab.lib.styles import ParagraphStyle
        from reportlab.platypus import HRFlowable, Image, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

        page_size = LETTER if template.page.size == "Letter" else A4
        doc = SimpleDocTemplate(
            str(target),
            pagesize=page_size,
            topMargin=_mm_to_points(template.page.margins.topMm),
            rightMargin=_mm_to_points(template.page.margins.rightMm),
            bottomMargin=_mm_to_points(template.page.margins.bottomMm),
            leftMargin=_mm_to_points(template.page.margins.leftMm),
        )
        story = []
        max_width = page_size[0] - doc.leftMargin - doc.rightMargin

        for block in _parse_blocks(markdown):
            if block.kind in {"heading", "paragraph", "list_item"}:
                prefix = block.prefix if block.kind == "list_item" else ""
                text = html.escape(prefix) + self._pdf_inline_markup(block.inlines or [], template, block)
                style = self._pdf_style(template, block)
                if block.kind == "list_item":
                    style.leftIndent = 18 * max(block.level, 1)
                if block.quote:
                    style.leftIndent += 18
                story.append(Paragraph(text or "&nbsp;", style))
            elif block.kind == "code":
                code_block = ExportBlock(kind="code")
                story.append(Paragraph(html.escape(block.text).replace("\n", "<br/>"), self._pdf_style(template, code_block)))
            elif block.kind == "hr":
                story.append(HRFlowable(width="100%", color=_reportlab_color(template.document.horizontalRuleColor), thickness=0.8, spaceBefore=6, spaceAfter=8))
            elif block.kind == "table":
                table = self._pdf_table(block, template)
                if table is not None:
                    story.append(table)
                    story.append(Spacer(1, 8))
            elif block.kind == "blank_line":
                story.append(Spacer(1, _blank_line_height(template)))
            for image_path in self._block_images(block, project_root, relative_path):
                try:
                    image = Image(str(image_path))
                    if image.drawWidth > max_width:
                        ratio = max_width / image.drawWidth
                        image.drawWidth = max_width
                        image.drawHeight *= ratio
                    story.append(image)
                    story.append(Spacer(1, 6))
                except Exception:
                    continue

        doc.build(story)

    def _pdf_style(self, template: ExportTemplateConfig, block: ExportBlock):
        from reportlab.lib.styles import ParagraphStyle

        style_data = _read_style(template, block)
        font_name = _pdf_font_name(style_data["fontFamily"], bold=block.kind == "heading" and _style_is_bold(style_data))
        font_size = style_data["fontSizePt"]
        return ParagraphStyle(
            name=f"{block.kind}-{block.level}",
            fontName=font_name,
            fontSize=font_size,
            leading=font_size * template.paragraph.lineSpacing,
            textColor=_reportlab_color(style_data["color"]),
            spaceAfter=template.paragraph.spaceAfterPt,
        )

    def _pdf_inline_markup(self, runs: list[InlineRun], template: ExportTemplateConfig, block: ExportBlock) -> str:
        parts: list[str] = []
        style = _read_style(template, block)
        for run in runs:
            if run.image:
                parts.append(html.escape(run.alt or run.image))
                continue
            text = html.escape(run.text).replace("\n", "<br/>")
            if run.code:
                text = f'<font face="{_pdf_font_name(template.code.fontFamily)}" color="{template.code.color}">{text}</font>'
            if run.bold:
                text = f"<b>{text}</b>"
            if run.italic:
                text = f"<i>{text}</i>"
            if run.underline:
                text = f"<u>{text}</u>"
            if block.kind == "heading" and _style_is_underlined(style):
                text = f"<u>{text}</u>"
            if run.strike:
                text = f"<strike>{text}</strike>"
            if run.link:
                text = f'<link href="{html.escape(run.link)}"><font color="{template.document.linkColor}"><u>{text}</u></font></link>'
            parts.append(text)
        return "".join(parts)

    def _pdf_table(self, block: ExportBlock, template: ExportTemplateConfig):
        from reportlab.platypus import Paragraph, Table, TableStyle

        rows = block.rows or []
        if not rows:
            return None
        table_data = []
        for row_index, row in enumerate(rows):
            cell_block = ExportBlock(kind="heading" if row_index == 0 else "paragraph", level=6)
            table_data.append([
                Paragraph(self._pdf_inline_markup(cell, template, cell_block), self._pdf_style(template, cell_block))
                for cell in row
            ])
        table = Table(table_data, hAlign="LEFT")
        table.setStyle(TableStyle([
            ("GRID", (0, 0), (-1, -1), 0.4, _reportlab_color("#E5E7EB")),
            ("BACKGROUND", (0, 0), (-1, 0), _reportlab_color("#FAFAFA")),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 6),
            ("RIGHTPADDING", (0, 0), (-1, -1), 6),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ]))
        return table

    def _block_images(self, block: ExportBlock, project_root: Path, relative_path: str) -> list[Path]:
        images: list[Path] = []
        for run in block.inlines or []:
            if run.image:
                resolved = _resolve_image_path(project_root, relative_path, run.image)
                if resolved is not None:
                    images.append(resolved)
        return images


export_service = ExportService()
